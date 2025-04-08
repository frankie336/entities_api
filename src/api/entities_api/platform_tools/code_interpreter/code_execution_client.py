import asyncio
import json
import os
from dataclasses import dataclass
from typing import Any, AsyncGenerator, Dict, Generator, Optional

import websockets
from dotenv import load_dotenv
from tenacity import retry, retry_if_exception_type, stop_after_attempt, wait_fixed
from websockets.legacy.protocol import WebSocketCommonProtocol

from entities_api.services.logging_service import LoggingUtility

load_dotenv()

logging_utility = LoggingUtility()


@dataclass
class ExecutionClientConfig:
    endpoint: str = os.getenv("CODE_EXECUTION_URL")
    timeout: float = 15.0
    retries: int = 3
    retry_delay: float = 2.0


class CodeExecutionClientError(Exception):
    pass


class ExecutionTimeoutError(CodeExecutionClientError):
    pass


class ExecutionSecurityViolation(CodeExecutionClientError):
    pass


class CodeExecutionClient:
    def __init__(self, config: Optional[ExecutionClientConfig] = None):
        self.config = config or ExecutionClientConfig()
        self._connection: Optional[WebSocketCommonProtocol] = None
        logging_utility.info(
            "CodeExecutionClient initialized with endpoint: %s", self.config.endpoint
        )

    async def __aenter__(self):
        await self.connect()
        return self

    async def __aexit__(self, exc_type, exc_val, exc_tb):
        await self.close()

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_fixed(2.0),
        retry=retry_if_exception_type((OSError, ConnectionRefusedError)),
    )
    async def connect(self):
        try:
            logging_utility.info(
                "Attempting to connect to WebSocket endpoint: %s", self.config.endpoint
            )
            self._connection = await asyncio.wait_for(
                websockets.connect(self.config.endpoint, ping_interval=None),
                timeout=self.config.timeout,
            )
            logging_utility.info("Successfully connected to WebSocket endpoint.")
        except asyncio.TimeoutError as e:
            logging_utility.error("Connection to %s timed out", self.config.endpoint)
            raise ExecutionTimeoutError(
                f"Connection to {self.config.endpoint} timed out"
            ) from e

    async def close(self):
        if self._connection:
            logging_utility.info("Closing WebSocket connection.")
            await self._connection.close()
            self._connection = None

    async def execute_code(
        self, code: str, **metadata: Dict[str, Any]
    ) -> AsyncGenerator[str, None]:
        if not self._connection:
            logging_utility.error("Execution attempt without active connection.")
            raise CodeExecutionClientError("Not connected to execution endpoint")

        try:
            logging_utility.info("Sending code execution request.")
            await self._connection.send(
                json.dumps({"code": code, "metadata": metadata})
            )

            while True:
                message = await self._connection.recv()
                if not message:
                    continue

                try:
                    data = json.loads(message)

                    if isinstance(data, dict):
                        if "error" in data:
                            logging_utility.error(
                                "ExecutionSecurityViolation: %s", data["error"]
                            )
                            raise ExecutionSecurityViolation(data["error"])

                        if "status" in data and "uploaded_files" in data:
                            # Final message with files
                            yield json.dumps(
                                {
                                    "type": "status",
                                    "content": data["status"],
                                    "execution_id": data.get("execution_id"),
                                    "uploaded_files": data.get("uploaded_files", []),
                                }
                            )
                            break
                        elif "status" in data:
                            # Intermediate status message
                            yield json.dumps(
                                {"type": "status", "content": data["status"]}
                            )
                            # Do NOT break — we may not be done yet
                        elif "output" in data:
                            yield json.dumps(
                                {"type": "output", "content": data["output"]}
                            )
                        else:
                            # Unrecognized dict – treat as raw
                            yield json.dumps(
                                {"type": "hot_code_output", "content": str(data)}
                            )
                    else:
                        yield json.dumps(
                            {"type": "hot_code_output", "content": str(data)}
                        )

                except json.JSONDecodeError:
                    logging_utility.warning("Received non-JSON output: %s", message)
                    yield json.dumps({"type": "hot_code_output", "content": message})

        except websockets.exceptions.ConnectionClosed:
            logging_utility.error("WebSocket connection closed unexpectedly.")
            raise CodeExecutionClientError("Connection closed unexpectedly")


class StreamOutput:
    def stream_output(self, code: str) -> Generator[str, None, None]:
        """Synchronous generator bridge for async execution"""
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)

        async def _async_wrapper() -> AsyncGenerator[str, None]:
            async with CodeExecutionClient() as client:
                async for chunk in client.execute_code(code):
                    yield chunk

        async_gen = _async_wrapper()
        try:
            while True:
                yield loop.run_until_complete(async_gen.__anext__())
        except StopAsyncIteration:
            pass


if __name__ == "__main__":
    logging_utility.info("Starting code execution.")
    output = StreamOutput()

    test_script = """
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document

# Pandas CSV output
df = pd.DataFrame({
    'Year': [2021, 2022, 2023],
    'Revenue': [100, 150, 200]
})
df.to_csv('report.csv', index=False)

# Matplotlib plot
plt.plot(df['Year'], df['Revenue'], marker='o')
plt.title('Revenue by Year')
plt.xlabel('Year')
plt.ylabel('Revenue')
plt.grid(True)
plt.savefig('plot.png')

# Word document
doc = Document()
doc.add_heading('Automated Report', level=1)
doc.add_paragraph('This report contains a CSV file, a chart, and a Word document.')
doc.save('summary.docx')
"""

    for chunk in output.stream_output(test_script):
        message = json.loads(chunk)
        if message["type"] == "status" and "uploaded_files" in message:
            logging_utility.info("Files: %s", message["uploaded_files"])
        else:
            logging_utility.info("Chunk: %s", message)
